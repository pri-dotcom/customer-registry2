import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PROJECT PLANNING & SCHEDULING\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Project Planning - 5 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # 1. Product Backlog & Sprint Schedule
    doc_builder.add_heading_1(doc, "1. Product Backlog & Sprint Estimations")
    doc_builder.add_paragraph(
        doc,
        "We estimated tasks using Fibonacci story points (1, 2, 3, 5, 8) representing complexity and effort. "
        "The project is structured into three Sprints."
    )
    
    headers_backlog = ["Sprint", "Functional Requirement (Epic)", "Story Number", "User Story / Task", "Story Points", "Priority", "Assigned Team Member"]
    data_backlog = [
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-01", "Technical Architecture", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-02", "ER Diagram", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-03", "Features", "1", "High", "Priyanka Degala (PD)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-04", "Roles and Responsibilities", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-05", "User Flow", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "TSK-06", "MVC Pattern", "1", "Medium", "Eswarasai Marre (EM)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "TSK-07", "Creating project folder", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "TSK-08", "Client Setup", "1", "High", "Priyanka Degala (PD)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "TSK-09", "Server Setup", "1", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "BACKEND DEVELOPMENT", "TSK-10", "Backend Structure", "2", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "BACKEND DEVELOPMENT", "TSK-11", "Development and Explanation", "8", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "TSK-12", "Configure MongoDB", "1", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "TSK-13", "Create Database Connection", "1", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "TSK-14", "Create Schema and Models", "3", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-3", "FRONTEND DEVELOPMENT", "TSK-15", "Frontend Structure", "2", "High", "Priyanka Degala (PD)"],
        ["Sprint-3", "FRONTEND DEVELOPMENT", "TSK-16", "Development and Explanation", "8", "High", "Priyanka Degala (PD)"],
        ["Sprint-3", "PROJECT EXECUTION", "TSK-17", "Steps for Project Execution", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-3", "PROJECT EXECUTION", "TSK-18", "Output Screenshots", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-3", "PROJECT EXECUTION", "TSK-19", "Drive Links", "1", "Medium", "Priyanka Degala (PD)"]
    ]
    doc_builder.add_styled_table(doc, headers_backlog, data_backlog, [0.8, 1.2, 0.8, 2.5, 0.8, 0.8, 1.1])
    
    # 2. Sprint Schedule Summary
    doc_builder.add_heading_2(doc, "1.2 Sprint Execution Schedule")
    headers_sprint = ["Sprint ID", "Total Story Points", "Duration", "Start Date", "End Date (Planned)", "Story Points Completed", "Actual Release Date"]
    data_sprint = [
        ["Sprint-1", "9", "6 Days", "10 July 2026", "15 July 2026", "9", "15 July 2026"],
        ["Sprint-2", "15", "6 Days", "17 July 2026", "22 July 2026", "15", "22 July 2026"],
        ["Sprint-3", "13", "6 Days", "24 July 2026", "29 July 2026", "13", "29 July 2026"]
    ]
    doc_builder.add_styled_table(doc, headers_sprint, data_sprint, [1.0, 1.2, 1.0, 1.2, 1.5, 1.5, 1.5])
    
    # 3. Gantt Chart Timeline
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "2. Project Timeline & Gantt Chart")
    doc_builder.add_paragraph(
        doc,
        "The project timeline spans 4 calendar weeks, divided into five phases."
    )
    
    headers_timeline = ["Phase ID", "Phase Name", "Tasks Covered", "Start Date", "End Date", "Dependencies"]
    data_timeline = [
        ["Phase-1", "Ideation & Requirements", "Problem Statements, Empathy Maps, DFDs.", "01 July 2026", "05 July 2026", "None"],
        ["Phase-2", "System Design", "Database Schemas, UML Diagramming, UI mockups.", "06 July 2026", "09 July 2026", "Phase-1"],
        ["Phase-3", "Core Development", "Sprint 1 & Sprint 2 implementation.", "10 July 2026", "22 July 2026", "Phase-2"],
        ["Phase-4", "Integration & Testing", "Sprint 3 features, Unit/Integration tests, UAT.", "23 July 2026", "29 July 2026", "Phase-3"],
        ["Phase-5", "Documentation", "Final reports compilation, video scripts, packaging.", "30 July 2026", "02 Aug 2026", "Phase-4"]
    ]
    doc_builder.add_styled_table(doc, headers_timeline, data_timeline, [0.8, 1.8, 2.5, 1.0, 1.0, 1.0])
    
    # 4. Project Milestones
    doc_builder.add_heading_1(doc, "3. Key Milestones")
    headers_milestones = ["Milestone ID", "Milestone Description", "Target Date", "Verification Criteria", "Status"]
    data_milestones = [
        ["MS-1", "Requirements Sign-off", "05 July 2026", "Functional requirements document approved.", "Completed"],
        ["MS-2", "Architecture Design Lock", "09 July 2026", "DB Schemas and UML diagrams verified.", "Completed"],
        ["MS-3", "Sprint 1 Complete", "15 July 2026", "Auth endpoints & customer profiles verified.", "Completed"],
        ["MS-4", "Sprint 2 Complete", "22 July 2026", "Complaint submissions and assignments verified.", "In Progress"],
        ["MS-5", "Sprint 3 Complete", "29 July 2026", "Custom fields and reports fully functional.", "Scheduled"],
        ["MS-6", "Final Deliverables Packaging", "02 Aug 2026", "All 10 project documents generated.", "Scheduled"]
    ]
    doc_builder.add_styled_table(doc, headers_milestones, data_milestones, [1.0, 2.0, 1.2, 2.3, 1.0])
    
    # 5. Risk Analysis & Mitigation
    doc_builder.add_heading_1(doc, "4. Risk Assessment & Mitigation Matrix")
    doc_builder.add_paragraph(
        doc,
        "We identified potential project risks and developed mitigation strategies to prevent project delays."
    )
    
    headers_risk = ["Risk Description", "Category", "Probability", "Impact", "Mitigation Strategy"]
    data_risk = [
        ["API latency spikes during ticket searches.", "Technical", "Low", "Medium", "Implement database indexes on key collections (customer, agent, status)."],
        ["Data loss during deployment or server restart.", "Database", "Low", "Critical", "Enable automated backups on MongoDB Atlas; write data using Mongoose transaction logs."],
        ["Scope creep during Sprint 3 implementation.", "Management", "Medium", "High", "Enforce strict MVP guidelines; defer advanced AI auto-routing to future scope."],
        ["Upload size limit causes server disk exhaustion.", "Infrastructure", "Medium", "Medium", "Limit attachment file size to 2MB using Multer validation rules."]
    ]
    doc_builder.add_styled_table(doc, headers_risk, data_risk, [1.8, 1.0, 1.0, 1.0, 2.7])
    
    # 6. Resource Allocation
    doc_builder.add_heading_1(doc, "5. Resource Allocation")
    headers_resources = ["Project Role", "Name / Resource", "Allocation", "Primary Area of Responsibility"]
    data_resources = [
        ["Team Lead & Frontend Developer", "Priyanka Degala (PD)", "100%", "Project Management, Client Setup, Frontend Structure, and React Development."],
        ["Backend & QA Developer", "Chandrakala Marri (CM)", "100%", "Technical Architecture, ER Diagram, Project Folder Creation, and QA/Testing."],
        ["Backend Developer", "Eswarasai Marre (EM)", "100%", "MVC Pattern design, Server Setup, and Backend API Development."],
        ["Database & Reports Integrator", "Triloka Pavani Mandhapati (TPM)", "100%", "MongoDB configuration, database connections, schemas/models definition, and reports integration."],
        ["Fullstack Developer", "Santhosh Nikhil Moka (SNM)", "100%", "User flow mapping, roles/responsibilities setup, project execution steps, and screenshots packaging."]
    ]
    doc_builder.add_styled_table(doc, headers_resources, data_resources, [1.5, 1.3, 1.0, 3.2])
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Project_Planning.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
