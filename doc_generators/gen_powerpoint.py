import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("POWERPOINT PRESENTATION GUIDE\n\n15-SLIDE PROJECT DECK OUTLINE")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Presentation Deck - 2 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # Overview
    doc_builder.add_heading_1(doc, "Overview of the Presentation")
    doc_builder.add_paragraph(
        doc,
        "This guide outlines the content for each slide of the PowerPoint presentation. "
        "Each section includes slide text, visual recommendations, and speaker notes."
    )
    
    slides = [
        ("Slide 1", "Title Slide: Customer Care Registry System", 
         "• Project Title: Customer Care Registry System (CCRS)\n• Subtitle: A MERN Stack Support Registry Portal\n• Platform: IBM SkillsBuild SmartBridge Internship\n• Team ID: SB-IBM-2026-07-0012\n• Team Members:\n  1. Priyanka Degala (Lead)\n  2. Chandrakala Marri\n  3. Eswarasai Marre\n  4. Triloka Pavani Mandhapati\n  5. Santhosh Nikhil Moka", 
         "Visual: Deep Navy Blue background with White text, centered logos of IBM SkillsBuild and SmartBridge.", 
         "Welcome everyone. Today, I am presenting our project, the Customer Care Registry System, developed during the IBM SkillsBuild SmartBridge Internship. This platform streamlines support ticket management for customers, agents, and administrators using the MERN stack."),
        
        ("Slide 2", "The Problem Statement", 
         "• Customer: No ticket tracking system, leading to long email waits.\n• Agent: Overloaded with spreadsheets, manually assigning tickets, and lacking details.\n• Admin: Fragmented logs, making SLA performance difficult to monitor.", 
         "Visual: Three-panel layout highlighting Customer, Agent, and Admin pain points.", 
         "Our problem identification focused on three key areas: customers lack tracking visibility, agents face workload bottlenecks from manual spreadsheets, and administrators lack real-time performance analytics."),
        
        ("Slide 3", "Project Objectives", 
         "• Centralize support tickets into a single secure database.\n• Automate status tracking and updates.\n• Streamline ticket assignments based on agent workloads.\n• Support dynamic custom fields for flexible form configurations.", 
         "Visual: Bulleted list with icons representing key objectives.", 
         "The project objectives are clear: centralize ticket logging, automate status tracking, balance agent workloads, and support dynamic custom fields to adapt forms to changing needs."),
        
        ("Slide 4", "The Solution Concept", 
         "• Dynamic Forms: Custom fields (text, number, date) defined by admin at runtime.\n• Communications Log: Text chat log built into each ticket details view.\n• Performance Evaluation: Automated customer feedback loops.", 
         "Visual: Features map showing the core functionality of the portal.", 
         "Our solution provides a role-based portal. Its main feature is custom field registration, allowing admins to modify forms at runtime without database migrations, accompanied by integrated chat logs."),
        
        ("Slide 5", "Solution Architecture", 
         "• Client Tier: React SPA client makes API calls using Axios.\n• Server Tier: Express/Node logic layer processes requests.\n• Data Tier: MongoDB Atlas cloud instance persists documents.", 
         "Visual: 3-Tier architecture diagram showing data flow between React, Node, and MongoDB.", 
         "This slide shows our 3-tier architecture. The React frontend communicates with the Express/Node API server, which performs validation and accesses the MongoDB Atlas cloud database."),
        
        ("Slide 6", "Technology Stack", 
         "• Frontend: React.js (Vite, Axios, CSS Modules)\n• Backend: Node.js, Express.js\n• Database: MongoDB Atlas (NoSQL)\n• Authentication: JWT (JSON Web Tokens), bcrypt.js", 
         "Visual: Grid diagram displaying the MERN stack logos and security components.", 
         "We selected the MERN stack for its scalability. We use Vite for fast React builds, Express for RESTful routing, MongoDB for flexible JSON schemas, and JWT for stateless auth."),
        
        ("Slide 7", "Customer Module & Workflows", 
         "• Register and manage profiles.\n• Raise complaints and upload file attachments.\n• Track ticket status via a visual timeline.\n• Submit ratings and feedback after resolution.", 
         "Visual: Flowchart outlining the customer registration, complaint, and feedback process.", 
         "For customers, the workflow starts with registration. They can then raise complaints with attachments, track status changes in real-time, and submit feedback ratings after resolution."),
        
        ("Slide 8", "Support Agent Module", 
         "• View assigned complaints queue.\n• Communicate with customers on tickets.\n• Transition statuses ('In Progress' to 'Resolved').\n• Write resolution details.", 
         "Visual: Layout of the Agent Dashboard, highlighting the assigned complaints list.", 
         "Support agents use a specialized view to monitor assigned tickets, communicate with customers, transition ticket statuses, and write resolution notes."),
        
        ("Slide 9", "Administrative Control Panel", 
         "• Manage customer and agent accounts.\n• Assign tickets to agents based on workloads.\n• Define dynamic custom fields for forms.\n• Monitor SLA metrics.", 
         "Visual: Layout of the Admin Dashboard, highlighting charts and assignment tables.", 
         "Administrators manage users, assign tickets to agents, define dynamic custom fields, and monitor SLA metrics via dashboards."),
        
        ("Slide 10", "Database Schema & Collections", 
         "• User Collections: Customers, Agents, Admins.\n• Complaint Collection: Title, description, status, assignedAgent ref.\n• Feedback Collection: Complaint, rating (1-5), comment.", 
         "Visual: Entity Relationship Diagram (ERD) mapping Mongoose collections.", 
         "This slide shows our database design. We use referenced document relations to connect customers, agents, complaints, and feedback records."),
        
        ("Slide 11", "Custom Fields & Notifications Engine", 
         "• Dynamic Custom Fields: Stored in a MongoDB Map object.\n• Notifications: In-app alerts logged in the database and sent on status updates.", 
         "Visual: Data flow diagram outlining how custom fields are rendered on forms.", 
         "Our schema uses MongoDB Map fields to support dynamic custom fields. System alerts trigger notifications on every status update."),
        
        ("Slide 12", "Testing & Quality Assurance", 
         "• Backend: Jest and Supertest unit and integration tests.\n• Functional: Tested all user scenarios (TC-001 to TC-012).\n• Performance: Evaluated latencies under load.", 
         "Visual: Results summary table showing a 100% pass rate for functional tests.", 
         "We verified the system using Jest for backend tests and functional testing for all user scenarios. Performance testing confirmed sub-200ms API response latencies."),
        
        ("Slide 13", "Project Results (Screenshots Overview)", 
         "• Visual layouts of the Home Page, login, and registration views.\n• Dashboards showing status timelines and chat logs.\n• Admin panel displaying analytics charts.", 
         "Visual: Carousel of system screenshot placeholders.", 
         "Our results demonstrate a fully functional portal. Key views include home, signup, customer status tracking timelines, and admin analytics dashboards."),
        
        ("Slide 14", "Future Scope", 
         "• AI Auto-Routing: Watson NLP to automatically assign tickets to agents.\n• WebSockets: Socket.io for real-time chat updates.\n• Cloud Storage: Move files to S3.", 
         "Visual: Bulleted list outlining future enhancement milestones.", 
         "Future enhancements will focus on AI auto-routing to assign tickets, Socket.io for real-time chat, and cloud storage (AWS S3) to make the server stateless."),
        
        ("Slide 15", "Conclusion & References", 
         "• CCRS delivers a secure, centralized ticket management portal.\n• The MERN stack provides a responsive, scalable architecture.\n• Project links: GitHub repo and Demo Video script.", 
         "Visual: Thank You text with the presenter's contact info and GitHub link.", 
         "In conclusion, the CCRS project meets all functional requirements, delivering a responsive customer care portal. Thank you for your time. I am happy to answer any questions.")
    ]
    
    for slide_id, title, bullets, visual, script in slides:
        doc_builder.add_heading_1(doc, f"{slide_id}: {title}")
        doc_builder.add_heading_2(doc, "Slide Content Outline:")
        doc_builder.add_paragraph(doc, bullets)
        
        doc_builder.add_heading_3(doc, "Visual Recommendation:")
        doc_builder.add_paragraph(doc, visual)
        
        doc_builder.add_callout(doc, "Presenter Speaker Notes", script, style_type="INFO")
        doc_builder.add_paragraph(doc, "") # Spacing
        
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "PowerPoint_Presentation.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
