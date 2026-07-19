import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("IDEATION PHASE DOCUMENT\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Ideation Phase - 4 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # 1. Problem Statement
    doc_builder.add_heading_1(doc, "1. Defined Problem Statements")
    doc_builder.add_paragraph(
        doc,
        "A well-articulated customer problem statement allows the design team to focus on resolving genuine operational issues "
        "faced by the stakeholders of the Customer Care Registry System (CCRS)."
    )
    
    headers_ps = ["Problem ID", "Statement Description", "Target Persona"]
    data_ps = [
        ["PS-1", "I am a Customer. I'm trying to report a service outage and track its resolution, but my only contact is a generic email address. Because there is no ticket tracking or updates, I feel ignored, frustrated, and inclined to switch service providers.", "Customer"],
        ["PS-2", "I am a Support Agent. I'm trying to resolve complaints quickly, but tickets are manually distributed via spreadsheets without clear prioritization. Because I cannot message customers directly on specific tickets, I waste time coordinating with internal teams, which makes me feel inefficient and overwhelmed.", "Support Agent"],
        ["PS-3", "I am a Customer Care Manager. I'm trying to monitor support performance and SLA compliance, but support logs are scattered across multiple data sheets. Because I lack centralized analytics or rating trends, I cannot identify bottlenecks, which makes me feel out of control.", "Administrator / Manager"]
    ]
    doc_builder.add_styled_table(doc, headers_ps, data_ps, [1.0, 4.5, 1.5])
    
    # 2. Empathy Map Canvas
    doc_builder.add_heading_1(doc, "2. Empathy Map Canvas")
    doc_builder.add_paragraph(
        doc,
        "The Empathy Map Canvas helps us understand our primary user's (Customer's) behavior, thoughts, actions, and feelings."
    )
    
    headers_em = ["Category", "User Observations & Expressions", "System Requirements Derived"]
    data_em = [
        ["Says", "'Why is there no update on my ticket?' 'I have to repeat my issue to every agent.'", "Provide a live timeline tracker; assign a single dedicated agent per ticket."],
        ["Thinks", "'My ticket was lost.' 'They don't care about my business.'", "Send automated email and in-app alerts on status updates."],
        ["Does", "Sends follow-up emails; searches for support numbers; writes negative reviews.", "Provide a chat log inside the complaint details view."],
        ["Feels", "Frustrated, ignored, anxious about service downtime.", "Use a clean dashboard layout that clearly shows active steps."],
        ["Pains", "Long delays, lack of transparency, security fears.", "Strict authentication checks, structured status progression, secure logs."],
        ["Gains", "Fast resolution, clear tracking, personalized updates.", "Admin-defined custom fields, instant notifications, rating feedback."]
    ]
    doc_builder.add_styled_table(doc, headers_em, data_em, [1.2, 3.3, 2.5])
    
    # 3. Brainstorming & Idea Prioritization
    doc_builder.add_heading_1(doc, "3. Brainstorming & Idea Prioritization")
    doc_builder.add_paragraph(
        doc,
        "Our team brainstormed a wide list of potential features. We grouped these features "
        "and prioritized them using an Impact-vs-Effort grid to establish the core product MVP."
    )
    
    headers_prior = ["Proposed Feature", "Category", "User Impact", "Development Effort", "Priority Status"]
    data_prior = [
        ["JWT Authentication & RBAC", "Security", "High", "Medium", "Core MVP (Must Have)"],
        ["Dynamic Custom Fields", "Flexibility", "High", "High", "Core MVP (Must Have)"],
        ["Interactive Analytics Dashboard", "Reporting", "High", "High", "Core MVP (Must Have)"],
        ["Real-time In-App Messaging", "Communication", "High", "High", "Core MVP (Must Have)"],
        ["External Email Server (SMTP)", "Notification", "Medium", "Medium", "Sprint-2 Integration"],
        ["SMS Gateway Integration", "Notification", "Medium", "High", "Future Scope"],
        ["AI Auto-Assignment Algorithm", "Intelligence", "High", "High", "Future Scope"],
        ["PDF Report Exporting", "Reporting", "Medium", "Medium", "Sprint-3 Integration"]
    ]
    doc_builder.add_styled_table(doc, headers_prior, data_prior, [2.0, 1.2, 1.0, 1.1, 1.2])
    
    # 4. Value Proposition Canvas
    doc_builder.add_heading_1(doc, "4. Value Proposition Canvas")
    doc_builder.add_paragraph(
        doc,
        "The Value Proposition Canvas aligns customer needs (Customer Profile) with our system's offerings (Value Map)."
    )
    doc_builder.add_heading_2(doc, "4.1 Customer Profile")
    doc_builder.add_bullet_point(doc, "Submit complaints, track ticket progress, communicate details, update profile.", bold_prefix="Customer Jobs: ")
    doc_builder.add_bullet_point(doc, "Long response delays, lack of updates, repeating explanations, data security fears.", bold_prefix="Customer Pains: ")
    doc_builder.add_bullet_point(doc, "Fast resolution, clear updates, personalized service, easy feedback options.", bold_prefix="Customer Gains: ")
    
    doc_builder.add_heading_2(doc, "4.2 Value Map")
    doc_builder.add_bullet_point(doc, "Centralized registry system with customer, agent, and admin interfaces.", bold_prefix="Products & Services: ")
    doc_builder.add_bullet_point(doc, "Live timeline tracker, direct chat messaging, role-based JWT security.", bold_prefix="Pain Relievers: ")
    doc_builder.add_bullet_point(doc, "Dynamic custom fields, ratings feedback, administrative analytics charts.", bold_prefix="Gain Creators: ")
    
    # 5. SWOT Analysis
    doc_builder.add_heading_1(doc, "5. SWOT Analysis")
    doc_builder.add_paragraph(
        doc,
        "SWOT Analysis evaluates internal Strengths and Weaknesses, alongside external Opportunities and Threats."
    )
    
    headers_swot = ["SWOT Dimension", "Factors Evaluated", "Strategic Action Plan"]
    data_swot = [
        ["STRENGTHS (Internal)", "MERN stack scalability, dynamic custom fields, token-based security.", "Leverage schema flexibility to implement fast customizations without downtime."],
        ["WEAKNESSES (Internal)", "Local file storage, lack of real-time WebSockets (relies on API polls).", "Refactor storage to S3; plan WebSocket integration in future sprints."],
        ["OPPORTUNITIES (External)", "Integration of AI chat assistance, auto-routing algorithms.", "Add AI auto-assignment to speed up ticket processing."],
        ["THREATS (External)", "Dependency on internet connectivity; potential database server latency.", "Implement database indexing; use redundant servers for high availability."]
    ]
    doc_builder.add_styled_table(doc, headers_swot, data_swot, [1.5, 2.5, 2.5])
    
    # 6. Root Cause Analysis
    doc_builder.add_heading_1(doc, "6. Root Cause Analysis")
    doc_builder.add_paragraph(
        doc,
        "We performed a Root Cause Analysis using the 5 Whys technique to identify the primary causes of slow complaint resolution."
    )
    doc_builder.add_paragraph(doc, "Problem Statement: Customer complaint resolution times exceed 5 days.")
    doc_builder.add_numbered_point(doc, "Why? Support agents are unaware of newly submitted tickets.", num_str="1. ")
    doc_builder.add_numbered_point(doc, "Why? Tickets are compiled manually in spreadsheets, causing delays.", num_str="2. ")
    doc_builder.add_numbered_point(doc, "Why? Scattered emails and phone calls are the only reporting channels.", num_str="3. ")
    doc_builder.add_numbered_point(doc, "Why? There is no centralized system to record and track tickets.", num_str="4. ")
    doc_builder.add_numbered_point(doc, "Why? The organization lacks a dedicated, integrated Customer Care Registry System (Root Cause).", num_str="5. ")
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Ideation_Phase.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
