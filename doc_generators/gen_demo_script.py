import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("DEMO VIDEO SCRIPT\n\n5-MINUTE NARRATION & VISUAL TIMELINE")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Video Script - 2 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # Intro
    doc_builder.add_heading_1(doc, "1. Overview of the Demo Video")
    doc_builder.add_paragraph(
        doc,
        "This document contains a structured, two-column script for a 5-minute video demonstration of the "
        "Customer Care Registry System. The left column details what is displayed on screen, and the right "
        "column contains the narration script."
    )
    
    headers_script = ["Time & Visual Cue", "Narration Voiceover Script"]
    data_script = [
        ["0:00 - 0:25\nVisual: Home Page\nScreen shows the CCRS Landing Page, introducing its features and navigation.", 
         "Hello and welcome. In this video, we demonstrate the Customer Care Registry System, a MERN stack application designed to streamline customer support operations. We start on the home landing page, which features clear login and register options for customers, agents, and administrators."],
        
        ["0:25 - 0:50\nVisual: Signup Page\nUser navigates to Register page, fills name, email, password, and phone number, then submits.", 
         "First, we will register a new customer account. We fill out the registration form. Upon submission, the system hashes the password using bcrypt, saves the record in MongoDB, and returns a secure JWT session token, redirecting the user to their dashboard."],
        
        ["0:50 - 1:15\nVisual: Customer Profile\nCustomer opens Profile page, uploads a profile image, updates their billing address, and submits.", 
         "On the Profile page, the customer can manage their personal details. We upload a profile picture and add a billing address. These details update in MongoDB, keeping the user profile current."],
        
        ["1:15 - 1:45\nVisual: Raise Complaint\nCustomer opens the Raise Complaint form, enters details, selects category and priority, and uploads a file.", 
         "Now, let's submit a support ticket. We open the Raise Complaint form and enter a title, description, category, and priority level. The form also supports custom fields configured by the admin. We click submit, creating a ticket in the database with status set to 'Pending'."],
        
        ["1:45 - 2:10\nVisual: Status Tracker\nCustomer opens their dashboard, showing a list of tickets and a visual status timeline tracker.", 
         "The customer is redirected to their dashboard, where they can monitor the progress of their tickets. We see our newly created ticket listed, with a status timeline tracking its progress in real-time."],
        
        ["2:10 - 2:40\nVisual: Admin Login & Panel\nUser logs out of the customer account, logs into the Admin account, and opens the Admin Dashboard.", 
         "Next, we switch to the administrator view. We log out and log back in using admin credentials. The Admin Dashboard displays statistics for total pending, resolved, and closed tickets, alongside analytics charts."],
        
        ["2:40 - 3:10\nVisual: User Management & Fields\nAdmin navigates to User Management, reviews customer and agent lists, then configures a custom field.", 
         "The admin has full visibility over the system. Under User Management, we can review customer and agent lists. Under Custom Fields, we can create new registration fields, which will render dynamically on customer forms."],
        
        ["3:10 - 3:40\nVisual: Ticket Assignment\nAdmin opens the unassigned ticket list, views agent workloads, selects an agent, and assigns the ticket.", 
         "Let's assign our customer's ticket. The admin opens the unassigned ticket list, checks agent workloads to balance assignments, select an agent, and click assign. The ticket status updates to 'Assigned' and alerts the agent."],
        
        ["3:40 - 4:10\nVisual: Agent Dashboard\nUser logs into the Agent account. The dashboard shows assigned ticket counts and pending tasks.", 
         "Now, let's log in as the assigned support agent. The Agent Dashboard displays assigned ticket counts and pending tasks, ensuring clear visibility over their workload."],
        
        ["4:10 - 4:30\nVisual: Ticket Details & Chat\nAgent opens the ticket, views details, changes status to 'In Progress', and sends a message.", 
         "The agent opens the ticket, views details, and changes the status to 'In Progress'. We can also use the integrated chat feature to send a message directly to the customer, keeping communication centralized."],
        
        ["4:30 - 4:45\nVisual: Resolution Submission\nAgent inputs resolution notes, sets status to 'Resolved', and submits the update.", 
         "Once the issue is resolved, the agent inputs resolution notes, sets the ticket status to 'Resolved', and submits. This logs the resolution timestamp and sends an alert to the customer."],
        
        ["4:45 - 5:00\nVisual: Customer Feedback\nUser logs back into the customer account, reviews resolution notes, and submits a 5-star rating.", 
         "Finally, we log back in as the customer. We open the resolved ticket, review the resolution notes, and submit a 5-star rating with feedback. The ticket status transitions to 'Closed', completing the support workflow. Thank you for watching."]
    ]
    doc_builder.add_styled_table(doc, headers_script, data_script, [2.2, 4.8])
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Demo_Video_Script.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
