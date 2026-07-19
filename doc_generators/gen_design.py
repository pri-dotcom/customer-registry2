import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PROJECT DESIGN DOCUMENT\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Project Design - 4 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # 1. Problem - Solution Fit
    doc_builder.add_heading_1(doc, "1. Problem - Solution Fit")
    doc_builder.add_paragraph(
        doc,
        "The Customer Care Registry System aligns features with the core requirements of its users. "
        "The following table maps user problems directly to features and their operational impacts."
    )
    
    headers_fit = ["Documented Pain Point", "Proposed Feature Solution", "Mechanisms of Resolution"]
    data_fit = [
        ["Customer lacks visibility into ticket progress.", "Visual Status Timeline & Live Tracker", "Displays active ticket stage (Pending -> Assigned -> In Progress -> Resolved -> Closed) in real-time."],
        ["Manual ticket assignment leads to agent fatigue.", "Admin Assignment Module with Agent Load Counter", "Admins see active ticket loads before assignment, avoiding overloading and balancing distribution."],
        ["Inquiries require varying inputs depending on department.", "Dynamic Custom Fields (text, date, number)", "Admins define custom fields globally. The system renders them dynamically on forms, matching specific needs."],
        ["Lack of clear analytics limits resource allocation.", "Admin Interactive Dashboard & Charts", "Displays live ticket counts, average resolution times, and categories using Chart.js graphs."]
    ]
    doc_builder.add_styled_table(doc, headers_fit, data_fit, [2.0, 2.0, 2.5])
    
    # 2. Proposed Solution
    doc_builder.add_heading_1(doc, "2. Proposed Solution")
    headers_sol = ["Parameter", "Description & Implementation Details"]
    data_sol = [
        ["Problem Statement", "Slow resolution times and lack of status tracking in customer support operations."],
        ["Solution Description", "MERN Stack portal with role-based JWT access, dynamic forms, chat logs, and admin analytics."],
        ["Novelty / Uniqueness", "Admin-defined custom fields stored as MongoDB Maps allow forms to be modified at runtime without database migrations."],
        ["Social Impact", "Reduces support delays in utility services (telecom, internet, energy), leading to higher user satisfaction."],
        ["Revenue Model", "SaaS subscription tiers based on monthly ticket volumes or active support agent licenses."],
        ["Scalability", "Stateless API server containers allow horizontal scalability; MongoDB indexes prevent query delays."]
    ]
    doc_builder.add_styled_table(doc, headers_sol, data_sol, [1.8, 4.7])
    
    # 3. Solution Architecture
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "3. Solution Architecture")
    doc_builder.add_paragraph(
        doc,
        "The CCRS implements a 3-Tier Client-Server Architecture. In this design, "
        "the Presentation Layer (React SPA client) is separated from the Application Layer (Express/Node API server) and the "
        "Data Layer (MongoDB Atlas NoSQL cloud instance)."
    )
    doc_builder.add_bullet_point(
        doc,
        "Built with React and Vite. It renders views, coordinates internal routing via React Router, "
        "and handles API calls via Axios. Users see a customized interface depending on their JWT role (Customer, Agent, Admin).",
        bold_prefix="Presentation Layer (React): "
    )
    doc_builder.add_bullet_point(
        doc,
        "Run on Node.js using Express.js. It defines API endpoints, validates JWTs in auth middleware, "
        "executes logic in controllers (like auth, complaints, and feedback), and returns JSON payloads.",
        bold_prefix="Logic Layer (Node/Express): "
    )
    doc_builder.add_bullet_point(
        doc,
        "A MongoDB Atlas cloud database. Mongoose schemas validate models and manage collection structures.",
        bold_prefix="Persistence Layer (MongoDB): "
    )
    
    # 4. Database Design (Collections Schemas)
    doc_builder.add_heading_1(doc, "4. Database Design & Schemas")
    doc_builder.add_paragraph(
        doc,
        "The data design is schema-enforced via Mongoose. The system consists of 7 primary collections. "
        "Below are the field-level schemas."
    )
    
    # Customer Schema
    doc_builder.add_heading_2(doc, "4.1 Customer Schema (customers)")
    headers_cust = ["Field Name", "Type", "Constraints", "Description"]
    data_cust = [
        ["_id", "ObjectId", "Primary Key, Auto-gen", "Unique system identifier for the customer record."],
        ["name", "String", "Required, Trimmed", "Customer's full name."],
        ["email", "String", "Required, Unique, Lowercase", "Authentication email address."],
        ["password", "String", "Required, Min length 6", "Bcrypt hashed password string."],
        ["phone", "String", "Required, Trimmed", "Contact phone number."],
        ["role", "String", "Required, Default='customer'", "System access role identifier."],
        ["profileImage", "String", "Default=''", "Relative URL path to the uploaded profile picture."],
        ["address", "String", "Default=''", "Optional billing/contact address details."],
        ["customFields", "Map of String", "Default={}", "Dynamic key-value pairs defined by admin custom fields."],
        ["isActive", "Boolean", "Default=true", "Identifies if the account is active or deactivated."]
    ]
    doc_builder.add_styled_table(doc, headers_cust, data_cust, [1.2, 1.2, 1.6, 2.5])
    
    # Complaint Schema
    doc_builder.add_heading_2(doc, "4.2 Complaint Schema (complaints)")
    headers_comp = ["Field Name", "Type", "Constraints", "Description"]
    data_comp = [
        ["_id", "ObjectId", "Primary Key, Auto-gen", "Unique complaint ticket identifier."],
        ["customer", "ObjectId", "Required, Ref='Customer'", "References the customer who raised the complaint."],
        ["assignedAgent", "ObjectId", "Ref='Agent', Default=null", "References the support agent assigned to resolve the issue."],
        ["title", "String", "Required, Trimmed", "Short summary title of the complaint."],
        ["description", "String", "Required, Trimmed", "Detailed description of the customer's issue."],
        ["category", "String", "Required", "Issue category classification (e.g., Broadband, Billing)."],
        ["priority", "String", "Enum: Low/Medium/High/Critical", "SLA priority level of the complaint."],
        ["status", "String", "Enum: Pending/Assigned/In Progress/Resolved/Closed", "Current status of the complaint workflow."],
        ["resolution", "String", "Default=''", "Notes written by the agent upon resolving the issue."],
        ["resolvedAt", "Date", "Default=null", "Timestamp when the ticket was resolved."],
        ["attachments", "Array of Objects", "Default=[]", "Lists attachment filenames and storage URLs."]
    ]
    doc_builder.add_styled_table(doc, headers_comp, data_comp, [1.2, 1.2, 1.6, 2.5])
    
    # 5. UML Diagrams Description
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "5. UML Diagrams Mapping")
    doc_builder.add_paragraph(
        doc,
        "To illustrate the runtime execution and architectural layout, we describe the system's structural "
        "and behavioral diagrams."
    )
    
    doc_builder.add_heading_2(doc, "5.1 Use Case Diagram")
    doc_builder.add_paragraph(
        doc,
        "The Use Case Diagram defines interactions between external actors and CCRS use cases: "
        "- Customers interact with: Register Account, Log In, Manage Profile, Raise Complaint, Track Status, Submit Feedback. "
        "- Support Agents interact with: Log In, View Assigned Tickets, Message Customer, Submit Resolution, Update Profile. "
        "- Administrators interact with: Log In, Manage Customers & Agents, Assign Tickets, Configure Custom Fields, View Analytics."
    )
    
    doc_builder.add_heading_2(doc, "5.2 Sequence Diagram (Raise & Assign Complaint)")
    doc_builder.add_paragraph(
        doc,
        "The Sequence Diagram maps chronological lifecycle communications: "
        "1. Customer: Submits complaint details via the React UI. "
        "2. React Client: Validates form details and sends a POST request with JWT auth headers to the API server. "
        "3. API Server: Validates the token, parses inputs, and calls `Complaint.create()` via Mongoose. "
        "4. Database: Writes the record and returns a success response. "
        "5. API Server: Triggers a notification and returns a 201 Created payload to the client, updating the dashboard. "
        "6. Admin: Selects the ticket on the Admin UI and assigns it to an agent. "
        "7. API Server: Updates the `assignedAgent` field, updates status to 'Assigned', saves to the DB, and alerts the agent."
    )
    
    doc_builder.add_heading_2(doc, "5.3 Activity Diagram (Resolution Flow)")
    doc_builder.add_paragraph(
        doc,
        "The Activity Diagram maps decision logic: "
        "- Starts when an agent opens an assigned ticket. "
        "- The agent transitions the status to 'In Progress' and starts investigation. "
        "- Decision: Is more customer info needed? "
        "  - Yes: The agent posts a message, transitioning the state to await response. "
        "  - No: The agent resolves the technical issue, enters resolution notes, and saves the updates. "
        "- The status transitions to 'Resolved', and a notification is sent to the customer. "
        "- The customer logs in, reviews the resolution notes, and submits feedback. "
        "- The status transitions to 'Closed', completing the activity loop."
    )
    
    doc_builder.add_heading_2(doc, "5.4 Deployment Diagram")
    doc_builder.add_paragraph(
        doc,
        "The Deployment Diagram details physical hosting allocations: "
        "- Client Machine: Runs modern web browsers (Chrome, Edge, Safari) rendering the React Single Page Application (SPA). "
        "- Cloud Hosting Server: Runs a Node.js process managed by PM2 within a container, exposing port 5001 behind an Nginx reverse proxy. "
        "- Database Server Cluster: A secure MongoDB Atlas NoSQL replica-set cluster with database access limited to the API server's IP address."
    )
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Project_Design.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
