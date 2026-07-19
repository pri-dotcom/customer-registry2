import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("REQUIREMENT ANALYSIS DOCUMENT\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Requirement Analysis - 4 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # 1. Actors, Constraints and Assumptions
    doc_builder.add_heading_1(doc, "1. System Boundary and Stakeholders")
    
    doc_builder.add_heading_2(doc, "1.1 System Actors")
    doc_builder.add_bullet_point(doc, "Can register accounts, raise complaints, track statuses, exchange messages, and submit feedback.", bold_prefix="Customer: ")
    doc_builder.add_bullet_point(doc, "Can log in, view assigned complaints, communicate with customers, and submit resolutions.", bold_prefix="Support Agent: ")
    doc_builder.add_bullet_point(doc, "Can manage all users, assign tickets to agents, create custom registration fields, and view system reports.", bold_prefix="Administrator: ")
    
    doc_builder.add_heading_2(doc, "1.2 Constraints")
    doc_builder.add_bullet_point(doc, "The system requires an active internet connection to submit tickets or view updates.", level=0)
    doc_builder.add_bullet_point(doc, "File attachments are limited to 2MB per upload to control storage utilization.", level=0)
    
    doc_builder.add_heading_2(doc, "1.3 Assumptions")
    doc_builder.add_bullet_point(doc, "Users have basic web browser access and literacy.", level=0)
    doc_builder.add_bullet_point(doc, "Administrator verifies newly created agent accounts to ensure security.", level=0)
    
    # 2. Functional Requirements
    doc_builder.add_heading_1(doc, "2. Functional Requirements")
    headers_fr = ["FR No.", "Epic / Module", "Sub-Requirement (Story / Sub-Task)"]
    data_fr = [
        ["FR-1", "User Registration", "Registration through form; Custom field integration; Phone validation."],
        ["FR-2", "User Login", "JWT Token generation; Role validation (Customer, Agent, Admin); Password verification."],
        ["FR-3", "Complaint Registration", "Submit ticket title, category, priority, description; File attachments; Status set to 'Pending'."],
        ["FR-4", "Complaint Assignment", "Admin interface showing workload; Manual selection dropdown; Agent allocation update."],
        ["FR-5", "Feedback loops", "Customer ratings out of 5 stars; Text comments; Only allowed for 'Resolved' status tickets."],
        ["FR-6", "Reports & Charts", "Stats tracking; Interactive bar/pie charts displaying ticket load by category."]
    ]
    doc_builder.add_styled_table(doc, headers_fr, data_fr, [1.0, 2.0, 4.0])
    
    # 3. Non-Functional Requirements
    doc_builder.add_heading_1(doc, "3. Non-Functional Requirements")
    headers_nfr = ["NFR No.", "Non-Functional Requirement", "Description"]
    data_nfr = [
        ["NFR-1", "Usability", "Clean, responsive user interface supporting mobile and desktop layouts."],
        ["NFR-2", "Security", "JWT sessions; password hashing using bcrypt; CORS header protections."],
        ["NFR-3", "Reliability", "Uptime rate of 99.9%; database replica sets prevent data loss during hardware failure."],
        ["NFR-4", "Performance", "API response time under 200ms; database query indexes on user and complaint IDs."],
        ["NFR-5", "Availability", "Deployment on PM2 clustering, supporting automatic process restarts."],
        ["NFR-6", "Scalability", "Stateless server design allows horizontal scalability under load balancers."]
    ]
    doc_builder.add_styled_table(doc, headers_nfr, data_nfr, [1.0, 2.2, 3.8])
    
    # 4. User Stories
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "4. User Stories & Acceptance Criteria")
    doc_builder.add_paragraph(
        doc,
        "Below is the complete Agile backlog containing user stories, priority estimates, and Sprint plans."
    )
    
    headers_us = ["User Type", "Epic", "Story No.", "User Story / Task", "Acceptance Criteria", "Priority", "Release"]
    data_us = [
        ["Customer", "Registration", "USN-1", "As a customer, I can register by entering my details.", "Saved in DB, password hashed.", "High", "Sprint-1"],
        ["Customer", "Login", "USN-2", "As a customer, I can log in using my email and password.", "Session JWT returned on success.", "High", "Sprint-1"],
        ["Customer", "Complaint", "USN-3", "As a customer, I can submit complaints with attachments.", "Status set to 'Pending'; files saved.", "High", "Sprint-2"],
        ["Customer", "Tracking", "USN-4", "As a customer, I can view status updates on my tickets.", "Status transitions shown on UI.", "Medium", "Sprint-2"],
        ["Admin", "Assignment", "USN-5", "As an admin, I can assign tickets to specific agents.", "Assigned agent saved; status='Assigned'.", "High", "Sprint-2"],
        ["Agent", "Dashboard", "USN-6", "As an agent, I can view my assigned complaints list.", "Only shows tickets matching agent ID.", "High", "Sprint-2"],
        ["Agent", "Resolution", "USN-7", "As an agent, I can write resolutions and close tickets.", "Status updated to 'Resolved'; notes saved.", "High", "Sprint-3"],
        ["Customer", "Feedback", "USN-8", "As a customer, I can rate the resolution out of 5 stars.", "Saved in Feedback collection.", "Medium", "Sprint-3"],
        ["Admin", "Analytics", "USN-9", "As an admin, I can view charts showing ticket statistics.", "Bar/pie charts render correctly.", "High", "Sprint-3"]
    ]
    doc_builder.add_styled_table(doc, headers_us, data_us, [1.0, 1.0, 0.8, 1.8, 1.2, 0.8, 0.8])
    
    # 5. Customer Journey Map
    doc_builder.add_heading_1(doc, "5. Customer Journey Map")
    headers_cjm = ["Stage", "Actions", "Thinking & Goals", "Touchpoints", "Pain Points", "Opportunities"]
    data_cjm = [
        ["1. Signup", "Fills signup form.", "Is my data secure?", "Registration View", "Generic forms.", "Dynamic verification."],
        ["2. Raise Ticket", "Enters details, submits.", "Will this get assigned fast?", "Complaint Form", "Categories missing.", "Admin-defined categories."],
        ["3. Monitoring", "Refreshes dashboard.", "Has anyone started working?", "Status Timeline", "No updates.", "Status transition emails."],
        ["4. Resolution", "Reviews agent notes.", "Is my issue fully fixed?", "Resolution Panel", "Notes too brief.", "Allow rejection option."],
        ["5. Feedback", "Clicks rating stars.", "I want to rate this agent.", "Feedback Panel", "Survey too long.", "Simple 5-star ranking."]
    ]
    doc_builder.add_styled_table(doc, headers_cjm, data_cjm, [1.0, 1.0, 1.2, 1.0, 1.1, 1.2])
    
    # 6. Use Case Description
    doc_builder.add_heading_1(doc, "6. Use Case Descriptions")
    doc_builder.add_heading_3(doc, "Use Case: Raise Complaint")
    doc_builder.add_paragraph(doc, "Logged-in Customer.", bold_prefix="Primary Actor: ")
    doc_builder.add_paragraph(doc, "Customer fills ticket details, selects category/priority, and submits the form.", bold_prefix="Basic Flow: ")
    doc_builder.add_paragraph(doc, "Database writes new Complaint document with status set to 'Pending'; triggers dashboard updates.", bold_prefix="Postcondition: ")
    
    # 7. Technology Stack
    doc_builder.add_heading_1(doc, "7. Technology Stack Specifications")
    
    doc_builder.add_heading_2(doc, "7.1 Components & Technologies")
    headers_t = ["S.No", "Component", "Description", "Technology"]
    data_t = [
        ["1", "User Interface", "Single Page Application frontend.", "React.js, HTML5, CSS3, Vite"],
        ["2", "Application Logic", "RESTful API Server routing and controllers.", "Node.js, Express.js"],
        ["3", "Database", "Document persistent data storage.", "MongoDB Atlas (NoSQL)"],
        ["4", "Authentication", "Stateless session tokens.", "JSON Web Tokens (JWT)"],
        ["5", "Security", "Cryptographic password hashing.", "bcrypt.js"],
        ["6", "File Storage", "Handling file attachments.", "Local Server Storage (via Multer)"]
    ]
    doc_builder.add_styled_table(doc, headers_t, data_t, [0.6, 1.8, 2.8, 1.8])
    
    doc_builder.add_heading_2(doc, "7.2 Application Characteristics")
    headers_char = ["S.No", "Characteristics", "Description", "Technology"]
    data_char = [
        ["1", "Open-Source Frameworks", "Use verified, modern open-source modules.", "Express.js, Mongoose, React, Axios, bcrypt.js, cors"],
        ["2", "Security Implementations", "Encrypt storage and prevent cross-origin leak vulnerabilities.", "Bcrypt password hashing, JWT stateless authorization, CORS middleware configurations"],
        ["3", "Scalable Architecture", "Utilize a modular, stateless 3-tier layout.", "3-Tier architecture (React SPA client, Node API, Mongo Atlas)"],
        ["4", "Availability", "Use automatic restart monitoring software.", "PM2 cluster management, Atlas replica-set deployment"],
        ["5", "Performance", "Support fast response latencies under concurrent access.", "MongoDB indexes, stateless session operations, async Event Loop processes"]
    ]
    doc_builder.add_styled_table(doc, headers_char, data_char, [0.6, 1.8, 2.8, 1.8])
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Requirement_Analysis.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
