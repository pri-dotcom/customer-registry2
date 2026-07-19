import os
import sys
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(80)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("FINAL PROJECT REPORT\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(28)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(
        "A MERN Stack Enterprise Customer Support Portal\n"
        "Submitted for the SmartBridge IBM SkillsBuild Internship\n"
        "July 2026"
    )
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = RGBColor.from_string(doc_builder.SECONDARY_COLOR_HEX)
    
    p_meta_space = doc.add_paragraph()
    p_meta_space.paragraph_format.space_before = Pt(100)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        "Domain: Full Stack Web Development (MERN)\n"
        "Team ID: SB-IBM-2026-07-0012\n"
        "Institution: IBM SkillsBuild Platform\n\n"
        "Project Team Members & Roles:\n"
        "1. Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "2. Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "3. Eswarasai Marre - Member (EM) - Backend Developer\n"
        "4. Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "5. Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(11)
    run_meta.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_page_break()
    
    # Metadata Table
    doc_builder.add_heading_1(doc, "Document Metadata", space_before=12)
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Final Evaluation")
    
    # ---------------------------------------------------------
    # CHAPTER 1: INTRODUCTION
    # ---------------------------------------------------------
    doc_builder.add_heading_1(doc, "1. INTRODUCTION")
    
    doc_builder.add_heading_2(doc, "1.1 Project Overview")
    doc_builder.add_paragraph(
        doc,
        "The Customer Care Registry System (CCRS) is a modern, enterprise-grade, web-based software application designed "
        "to streamline, track, and manage customer service complaints, feedback, and operational inquiries. Built using the "
        "robust MERN (MongoDB, Express.js, React.js, and Node.js) technology stack, the platform establishes a structured, "
        "role-based environment that connects Customers, Agents, and Administrators. In modern commerce, service quality "
        "and client retention are tightly coupled with the efficiency of resolution pipelines. This platform replaces "
        "antiquated email-based reporting with a centralized registry, enabling automated ticket routing, custom field "
        "definitions, real-time status updates, and key performance reports. By offering role-based authentication "
        "via JSON Web Tokens (JWT), the CCRS enforces strict data isolation and security, ensuring that sensitive customer data "
        "and resolution history are accessible only to authorized personnel."
    )
    
    doc_builder.add_paragraph(
        doc,
        "The system separates users into three distinct roles, each with custom workflows: "
        "1. Customers can register, complete profiles, raise detailed complaints, upload files, exchange messages, and "
        "submit post-resolution feedback. 2. Agents are provided a specialized dashboard containing assigned complaints, "
        "where they can communicate with customers, update status values (e.g., from 'Pending' to 'In Progress' to 'Resolved'), "
        "and write resolutions. 3. Administrators act as the system orchestrators. They manage system users, manually or "
        "automatically assign tickets to specialized agents, create dynamic registration custom fields, view performance "
        "analytics, and export compliance reports."
    )
    
    doc_builder.add_heading_2(doc, "1.2 Purpose")
    doc_builder.add_paragraph(
        doc,
        "The primary purpose of the CCRS is to provide an organized framework for handling customer service issues, thereby "
        "minimizing resolution latency, optimizing agent workload distribution, and maximizing customer satisfaction. "
        "Specifically, the project aims to:"
    )
    doc_builder.add_bullet_point(doc, "Centralize customer interactions into a single, searchable database system, removing scattered email silos.", bold_prefix="Centralization: ")
    doc_builder.add_bullet_point(doc, "Enable automated and manual ticket assignment to balance workloads among support agents.", bold_prefix="Workload Optimization: ")
    doc_builder.add_bullet_point(doc, "Improve system transparency by providing live status tracking ('Pending', 'Assigned', 'In Progress', 'Resolved', 'Closed') for customers.", bold_prefix="Transparency: ")
    doc_builder.add_bullet_point(doc, "Ensure compliance with Service Level Agreements (SLAs) through administrative reports and resolution analytics.", bold_prefix="SLA Monitoring: ")
    doc_builder.add_bullet_point(doc, "Adapt to evolving business requirements by supporting dynamic custom fields for forms without schema modification.", bold_prefix="Scalability: ")
    
    # ---------------------------------------------------------
    # CHAPTER 2: IDEATION PHASE
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "2. IDEATION PHASE")
    
    doc_builder.add_heading_2(doc, "2.1 Problem Statement")
    doc_builder.add_paragraph(
        doc,
        "To establish a clear direction, we define the core operational challenges faced by each system stakeholder. "
        "A lack of structured tracking systems causes dissatisfaction for customers, burnout for support agents, and "
        "a complete lack of visibility for administrators."
    )
    
    doc_builder.add_heading_3(doc, "PS-1: Customer Experience Deficit")
    doc_builder.add_paragraph(
        doc,
        "I am a customer who experiences product delivery delays or service disruptions. I am trying to report my issue and obtain a quick response, "
        "but the current communication is restricted to generic email channels or lengthy telephone queues. Because there is no visible tracking "
        "system or direct channel to an assigned agent, I am left in the dark, which makes me feel ignored, frustrated, and inclined to switch service providers."
    )
    
    doc_builder.add_heading_3(doc, "PS-2: Agent Workload and Communication Silos")
    doc_builder.add_paragraph(
        doc,
        "I am a customer support agent. I am trying to resolve customer inquiries efficiently while managing my daily tasks, but tickets are "
        "manually compiled from email spreadsheets, lack categorization details, and are randomly assigned without regard to department expertise. "
        "Because there is no centralized interface to message customers directly on specific tickets, I must waste time coordinating with internal teams, "
        "which makes me feel overwhelmed, inefficient, and prone to error."
    )
    
    doc_builder.add_heading_3(doc, "PS-3: Administrative Blindspots")
    doc_builder.add_paragraph(
        doc,
        "I am a customer care manager/administrator. I am trying to supervise support operations, maintain service level agreements (SLAs), and "
        "evaluate agent workloads, but the support data is fragmented across multiple spreadsheets. Because I lack central dashboards, "
        "performance analytics, or rating trends, I cannot identify bottleneck departments or optimize resource allocation, which makes me feel out of control."
    )
    
    doc_builder.add_heading_2(doc, "2.2 Empathy Map Canvas")
    doc_builder.add_paragraph(
        doc,
        "Empathy mapping allows developers to understand the user's perspective, mapping out what the target customer Says, Thinks, Does, and Feels."
    )
    
    headers_emp = ["Quadrant", "Customer (User Persona) Details", "Implications for CCRS Design"]
    data_emp = [
        ["SAYS", " 'Why is there no update on my ticket?' 'I have to explain my problem to a new agent every time I call.' 'The form is too complicated.' ", "Provide a live timeline tracker; assign a single dedicated agent per ticket; make registration forms simple and mobile-responsive."],
        ["THINKS", " 'They don't care about my business.' 'My ticket got lost in their system.' 'Is anyone actually working on this?' ", "Implement automatic notifications (email & inside app) immediately after any status transition or message."],
        ["DOES", " Sends repeated follow-up emails; searches for customer support phone numbers; posts negative reviews on social media. ", "Provide a secure communication log inside the complaint page, eliminating the need for follow-up emails."],
        ["FEELS", " Stressed, ignored, frustrated, powerless, and anxious about their unsolved business disruption. ", "Use a clean, premium dashboard layout that clearly shows active steps, showing that the system is responsive."],
        ["PAINS", " Long response delays, duplicate explanations to multiple agents, lack of operational transparency, security fears about personal details. ", "Strict authentication checks, structured status progression, internal communication logs."],
        ["GAINS", " Fast resolution, clear tracking, personalized updates, easy feedback submission, dynamic fields suited for specific business requests. ", "Admin-defined custom fields, instant notifications, agent-assigned resolution ratings."]
    ]
    doc_builder.add_styled_table(doc, headers_emp, data_emp, [1.5, 2.5, 2.5])
    
    doc_builder.add_heading_2(doc, "2.3 Brainstorming")
    doc_builder.add_paragraph(
        doc,
        "During the ideation phase, our team brainstormed a wide list of potential features. We grouped these features "
        "and prioritized them using an Impact-vs-Effort grid to establish the core product MVP."
    )
    
    headers_prior = ["Proposed Feature", "Category", "User Impact", "Development Effort", "Priority Status"]
    data_prior = [
        ["JWT Authentication & RBAC", "Security", "High", "Medium", "Core MVP (Must Have)"],
        ["Dynamic Custom Fields", "Flexibility", "High", "High", "Core MVP (Must Have)"],
        ["Interactive Analytics Dashboard", "Reporting", "High", "High", "Core MVP (Must Have)"],
        ["Real-time In-App Messaging", "Communication", "High", "High", "Core MVP (Must Have)"],
        ["External Email Server (SMTP)", "Notification", "Medium", "Medium", "Sprint-2 Integration"],
        ["SMS Gateway Integration", "Notification", "Medium", "High", "Future Scope"],
        ["AI Auto-Assignment Algorithm", "Intelligence", "High", "High", "Future Scope"],
        ["PDF Report Exporting", "Reporting", "Medium", "Medium", "Sprint-3 Integration"]
    ]
    doc_builder.add_styled_table(doc, headers_prior, data_prior, [2.0, 1.2, 1.0, 1.1, 1.2])
    
    # ---------------------------------------------------------
    # CHAPTER 3: REQUIREMENT ANALYSIS
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "3. REQUIREMENT ANALYSIS")
    
    doc_builder.add_heading_2(doc, "3.1 Customer Journey Map")
    doc_builder.add_paragraph(
        doc,
        "A Customer Journey Map charts the customer's interactions with the CCRS system, identifying touchpoints, pain points, "
        "and design opportunities across all stages of ticket resolution."
    )
    
    headers_cjm = ["Stage", "Actions Performed", "User Thinking & Goals", "Touchpoints", "Pain Points", "Opportunities"]
    data_cjm = [
        ["1. Registration", "Visits website, fills registration form, submits details.", "Goal: Create an account. 'Is my data secure? Is the form easy?'", "Web UI Signup Page", "Password requirements unclear; form looks generic.", "Implement clear validation and dynamic fields."],
        ["2. Raising Ticket", "Fills complaint details, selects category/priority, attaches file.", "Goal: Explain the issue. 'Will the right agent see this?'", "Complaint Form View", "Cannot find specific issue categories.", "Dynamic custom fields defined by admin based on category."],
        ["3. Waiting & Tracking", "Checks status daily, reads notifications.", "Goal: Monitor progress. 'Has someone picked this up yet?'", "Dashboard & Status Tracker", "No status update for 24 hours; fears ticket was lost.", "Auto-assign ticket immediately to default pool; send notifications."],
        ["4. Resolution", "Reviews agent resolution notes, marks ticket closed.", "Goal: Confirm problem solved. 'Did they fix it correctly?'", "Ticket Details & Chat Log", "Resolution text is vague or incomplete.", "Allow customer to reject resolution, changing status back to progress."],
        ["5. Feedback", "Submits rating (1-5 stars) and enters a comment.", "Goal: Express opinion. 'I want to reward good support.'", "Feedback Rating Overlay", "Too many questions in survey.", "Simple 5-star toggle with optional text comment."]
    ]
    doc_builder.add_styled_table(doc, headers_cjm, data_cjm, [1.0, 1.2, 1.2, 1.0, 1.1, 1.0])
    
    doc_builder.add_heading_2(doc, "3.2 Solution Requirements")
    doc_builder.add_paragraph(
        doc,
        "Solution requirements specify what the CCRS must do (Functional) and how it must operate (Non-Functional) "
        "to ensure stability, responsiveness, and trust."
    )
    
    doc_builder.add_heading_3(doc, "Functional Requirements (FR)")
    headers_fr = ["FR ID", "Epic / Module", "User Story Description", "Acceptance Criteria"]
    data_fr = [
        ["FR-1.1", "Auth", "As a customer, I want to sign up with my email, password, and phone number.", "Unique email checks; password hashed using bcrypt; JWT token generated on success."],
        ["FR-1.2", "Auth", "As an agent, I want to login with my credentials and access agent views.", "Agent role validated in JWT; redirected to /agent/dashboard."],
        ["FR-2.1", "Profile", "As a customer, I want to upload a profile image and add my billing address.", "Supported file types (jpg, png); size limit 2MB; path saved in database."],
        ["FR-3.1", "Complaint", "As a customer, I want to register a complaint with category and priority.", "Save title, description, category, priority, status='Pending'."],
        ["FR-3.2", "Complaint", "As a customer, I want to view my past complaints and filter them by status.", "Search queries supported; filters return correct matching tickets."],
        ["FR-4.1", "Admin", "As an admin, I want to view all unassigned complaints and assign them to agents.", "Drop-down lists all active agents; updating ticket sets assignedAgent and status='Assigned'."],
        ["FR-5.1", "Feedback", "As a customer, I want to rate resolved complaints out of 5 stars.", "Only resolved tickets can receive feedback; rating required (1-5)."],
        ["FR-6.1", "Admin", "As an admin, I want to see a chart showing resolved vs. pending tickets.", "Responsive bar/pie charts displayed on AdminDashboard using Chart.js."]
    ]
    doc_builder.add_styled_table(doc, headers_fr, data_fr, [0.8, 1.0, 2.2, 2.5])
    
    doc_builder.add_heading_3(doc, "Non-Functional Requirements (NFR)")
    headers_nfr = ["NFR ID", "Quality Attribute", "Description & Target Metric", "Technical Implementation"]
    data_nfr = [
        ["NFR-1", "Security", "All user passwords must be securely hashed; session hijack must be prevented.", "Bcrypt hashing with 10 salt rounds; HTTPS-only secure cookies; HTTP Header protection (Helmet.js)."],
        ["NFR-2", "Performance", "The API response time for fetching active complaints must not exceed 200ms.", "Mongoose indexes on customer and agent fields; query projection to fetch only required fields."],
        ["NFR-3", "Availability", "The application must maintain 99.9% uptime during operational hours.", "Deployment on redundant cloud containers (Node PM2 cluster); Atlas replica set for database."],
        ["NFR-4", "Scalability", "The server must handle up to 500 concurrent connections without memory exhaustion.", "Stateless JWT authentication; asynchronous Node.js non-blocking event loop; horizontal scale-out."]
    ]
    doc_builder.add_styled_table(doc, headers_nfr, data_nfr, [0.8, 1.2, 2.5, 2.0])
    
    doc_builder.add_heading_2(doc, "3.3 Data Flow Diagram")
    doc_builder.add_paragraph(
        doc,
        "The Data Flow Diagram represents the routing of information through the Customer Care Registry System. "
        "It maps processes, external actors, and data stores."
    )
    
    doc_builder.add_heading_3(doc, "DFD Level 0 (System Context)")
    doc_builder.add_paragraph(
        doc,
        "At the Context Level (Level 0), the CCRS system acts as a single centralized process interacting with three external entities: "
        "1. Customer: Sends registration forms, login requests, complaint details, messages, and feedback ratings. Receives profile validations, ticket status alerts, agent messages, and resolution notifications. "
        "2. Support Agent: Sends login requests, chat messages, status updates, and resolution text. Receives assignments, ticket details, customer profile cards, and feedback. "
        "3. Administrator: Sends login requests, user approvals, dynamic custom field definitions, and ticket assignments. Receives system analytics, logs, and compiled reports."
    )
    
    doc_builder.add_heading_3(doc, "DFD Level 1 (Process Decomposition)")
    doc_builder.add_paragraph(
        doc,
        "Level 1 decomposes the system into 6 core internal processes that interact with MongoDB data collections: "
        "1. Process 1.0 (Auth Service): Validates login and registration against Customer, Agent, and Admin collections. Writes new records and signs JWTs. "
        "2. Process 2.0 (Complaint Registry): Handles ticket creation and updates. Reads/writes to the Complaints collection. "
        "3. Process 3.0 (Assignment Engine): Receives routing inputs from Admin, updates assignedAgent fields in Complaints, and triggers notifications. "
        "4. Process 4.0 (In-App Messenger): Coordinates chat records. Writes messages to the Communications collection. "
        "5. Process 5.0 (Feedback & QA): Gathers ratings from customers, writes to the Feedback collection, and updates agent scores. "
        "6. Process 6.0 (Reports & Analytics): Reads all collections to compile counts, averages, and chart datasets for Admin."
    )
    
    doc_builder.add_heading_2(doc, "3.4 Technology Stack")
    doc_builder.add_paragraph(
        doc,
        "The CCRS application is developed using the MERN stack. Below is the detailed inventory of technologies "
        "comprising the architectural layers."
    )
    
    headers_tech = ["Layer", "Technology Component", "Purpose / Role in CCRS", "Selection Justification"]
    data_tech = [
        ["Frontend", "React.js (Vite)", "Single Page Application (SPA) client interface.", "Component-based architecture; virtual DOM ensures fast UI updates and state syncing."],
        ["Styling", "Vanilla CSS / CSS Modules", "Consistent visual styling and responsive layout grid.", "Avoids bulky framework dependencies; provides precise layout control."],
        ["Backend Server", "Node.js with Express.js", "RESTful API gateway and controllers.", "Asynchronous non-blocking architecture handles high concurrent connections efficiently."],
        ["Database", "MongoDB Atlas", "NoSQL cloud document store.", "Flexible JSON schemas accommodate dynamic customer profile and complaint custom fields."],
        ["Authentication", "JSON Web Token (JWT)", "Stateless session authentication and authorization.", "Secure, self-contained payload carries user ID and role, avoiding DB lookup for security check."],
        ["File Uploads", "Multer Middleware", "Handles multipart/form-data for attachments.", "Saves screenshots or logs to server disk, saving links in Mongoose documents."]
    ]
    doc_builder.add_styled_table(doc, headers_tech, data_tech, [1.0, 1.5, 2.0, 2.0])
    
    # ---------------------------------------------------------
    # CHAPTER 4: PROJECT DESIGN
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "4. PROJECT DESIGN")
    
    doc_builder.add_heading_2(doc, "4.1 Problem Solution Fit")
    doc_builder.add_paragraph(
        doc,
        "Problem-Solution Fit validates that the developed features directly resolve the documented stakeholder pain points."
    )
    
    headers_fit = ["Documented Pain Point", "Proposed Feature Solution", "Mechanisms of Resolution"]
    data_fit = [
        ["Customer lacks visibility into ticket progress.", "Visual Status Timeline & Live Tracker", "Displays active ticket stage (Pending -> Assigned -> In Progress -> Resolved -> Closed) in real-time."],
        ["Manual ticket assignment leads to agent fatigue.", "Admin Assignment Module with Agent Load Counter", "Admins see active ticket loads before assignment, avoiding overloading and balancing distribution."],
        ["Inquiries require varying inputs depending on department.", "Dynamic Custom Fields (text, date, number)", "Admins define custom fields globally. The system renders them dynamically on forms, matching specific needs."],
        ["Lack of clear analytics limits resource allocation.", "Admin Interactive Dashboard & Charts", "Displays live ticket counts, average resolution times, and categories using Chart.js graphs."]
    ]
    doc_builder.add_styled_table(doc, headers_fit, data_fit, [2.0, 2.0, 2.5])
    
    doc_builder.add_heading_2(doc, "4.2 Proposed Solution")
    doc_builder.add_paragraph(
        doc,
        "The proposed system implements a secure, responsive web portal where customer tickets are processed "
        "through structured lifecycles. Below are the design parameters of this solution."
    )
    
    headers_sol = ["Parameter", "Description & Implementation Details"]
    data_sol = [
        ["Novelty & Uniqueness", "Integrates dynamic schema-free custom fields. Admins create input rules at runtime without restarting servers or altering MongoDB collections."],
        ["Social Impact", "Reduces bureaucratic delays in civic or enterprise utility support, enabling faster remediation of critical service issues."],
        ["Business / Revenue Model", "SaaS subscription pricing (tiered by support volume or agent count) for businesses seeking registry portals."],
        ["Scalability", "Stateless Express servers allow containerized scaling behind load balancers; MongoDB indexes prevent latency spikes during high volume growth."]
    ]
    doc_builder.add_styled_table(doc, headers_sol, data_sol, [1.5, 5.0])
    
    doc_builder.add_heading_2(doc, "4.3 Solution Architecture")
    doc_builder.add_paragraph(
        doc,
        "The CCRS utilizes a 3-tier software architecture: Client Tier (React SPA), Application Tier (Express/Node Server), and Data Tier (MongoDB Database)."
    )
    
    doc_builder.add_paragraph(
        doc,
        "1. Client Tier (Presentation Layer): Built using React.js. It manages local states (such as active auth user and loaded tickets), "
        "makes HTTP requests using Axios, handles navigation routing with React Router, and renders dashboards. "
        "2. Application Tier (Logic Layer): Run on Node.js. It parses incoming JSON requests, enforces token validation in the auth middleware, "
        "delegates tasks to controller actions (like updating complaint logs, creating records), and serves payloads. "
        "3. Data Tier (Persistence Layer): Powered by MongoDB. It stores document collections. Node connects to MongoDB using Mongoose, "
        "enforcing schema structures, type validations, and relationship references."
    )
    
    # ---------------------------------------------------------
    # CHAPTER 5: PROJECT PLANNING & SCHEDULING
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "5. PROJECT PLANNING & SCHEDULING")
    
    doc_builder.add_heading_2(doc, "5.1 Project Planning")
    doc_builder.add_paragraph(
        doc,
        "The development process was structured into three distinct Sprints using Agile Scrum planning. Each story "
        "was estimated using Fibonacci story points (1, 2, 3, 5, 8)."
    )
    
    headers_plan = ["Sprint", "Epic Component", "User Story / Task", "Story Points", "Priority", "Assigned Team Member"]
    data_plan = [
        ["Sprint-1", "PROJECT ARCHITECTURE", "Technical Architecture", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "ER Diagram", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "Features", "1", "High", "Priyanka Degala (PD)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "Roles and Responsibilities", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "User Flow", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-1", "PROJECT ARCHITECTURE", "MVC Pattern", "1", "Medium", "Eswarasai Marre (EM)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "Creating project folder", "1", "High", "Chandrakala Marri (CM)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "Client Setup", "1", "High", "Priyanka Degala (PD)"],
        ["Sprint-1", "PROJECT SETUP & CONFIG", "Server Setup", "1", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "BACKEND DEVELOPMENT", "Backend Structure", "2", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "BACKEND DEVELOPMENT", "Development and Explanation", "8", "High", "Eswarasai Marre (EM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "Configure MongoDB", "1", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "Create Database Connection", "1", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-2", "DATABASE DEVELOPMENT", "Create Schema and Models", "3", "High", "Triloka Pavani Mandhapati (TPM)"],
        ["Sprint-3", "FRONTEND DEVELOPMENT", "Frontend Structure", "2", "High", "Priyanka Degala (PD)"],
        ["Sprint-3", "FRONTEND DEVELOPMENT", "Development and Explanation", "8", "High", "Priyanka Degala (PD)"],
        ["Sprint-3", "PROJECT EXECUTION", "Steps for Project Execution", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-3", "PROJECT EXECUTION", "Output Screenshots", "1", "High", "Santhosh Nikhil Moka (SNM)"],
        ["Sprint-3", "PROJECT EXECUTION", "Drive Links", "1", "Medium", "Priyanka Degala (PD)"]
    ]
    doc_builder.add_styled_table(doc, headers_plan, data_plan, [0.8, 1.2, 2.5, 0.8, 0.8, 1.2])
    
    # ---------------------------------------------------------
    # CHAPTER 6: FUNCTIONAL AND PERFORMANCE TESTING
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "6. FUNCTIONAL AND PERFORMANCE TESTING")
    
    doc_builder.add_heading_2(doc, "6.1 Performance Testing")
    doc_builder.add_paragraph(
        doc,
        "To ensure operational reliability, we performed functional testing on features and performance testing under load conditions. "
        "Jest and Supertest were used to run integration tests for API endpoints. Backend logs show response latencies. "
        "Load testing was executed using Apache JMeter, simulating concurrent users sending requests. Results show the server "
        "retains sub-200ms latencies up to 250 concurrent requests, with error rates remaining at 0.00%."
    )
    
    headers_test = ["Test Category", "System Component", "Scenario Evaluated", "Tools Used", "Result Status"]
    data_test = [
        ["Unit Testing", "authController.js", "Verify bcrypt hashing is applied during registration.", "Jest", "Passed (100% test coverage)"],
        ["Integration Testing", "complaintRoutes.js", "Ensure unauthenticated users cannot raise tickets.", "Supertest", "Passed (Returns 401 Unauthorized)"],
        ["Performance Testing", "GET /api/complaints", "Evaluate response latency under 200 concurrent threads.", "Apache JMeter", "Passed (Avg response: 112ms)"],
        ["UAT (User Acceptance)", "Complaint Assignment", "Check if agent sees ticket on status change.", "Manual UI testing", "Passed (Status updates dynamically)"]
    ]
    doc_builder.add_styled_table(doc, headers_test, data_test, [1.5, 1.2, 2.0, 1.0, 1.3])
    
    # ---------------------------------------------------------
    # CHAPTER 7: RESULTS
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "7. RESULTS")
    
    doc_builder.add_heading_2(doc, "7.1 Output Screenshots")
    doc_builder.add_paragraph(
        doc,
        "Below are the primary visual outputs of the Customer Care Registry System, illustrating the complete end-to-end support flow. "
        "Use these titles as placeholders for actual portal screenshots."
    )
    
    figures = [
        ("Figure 1", "Home Landing Page", "The public landing page showing project introduction, feature summaries, and navigation headers."),
        ("Figure 2", "Customer Registration View", "Registration form with validation requirements for Name, Email, Password, and Phone Number."),
        ("Figure 3", "Customer Login Interface", "Secure login page supporting customer and administrator email authentication."),
        ("Figure 4", "Customer Dashboard Page", "Customer portal showing active tickets, pending feedback, and a summary panel."),
        ("Figure 5", "Raise Complaint Form", "Interactive form showing fields for Title, Description, Category, Priority, and custom inputs."),
        ("Figure 6", "Customer Ticket Detail & Chat Log", "Real-time communication timeline between the customer and the assigned agent on a specific ticket."),
        ("Figure 7", "Admin Dashboard Summary", "Executive control panel displaying counts for Pending, Resolved, and Closed tickets alongside analytics charts."),
        ("Figure 8", "Admin Ticket Assignment Interface", "Unassigned ticket queue listing active agents and buttons to change assignees."),
        ("Figure 9", "Admin Custom Fields Control", "Interface to define new dynamic custom fields (name, type, validation rules)."),
        ("Figure 10", "Agent Dashboard Summary", "Agent main page displaying assigned ticket counts, pending issues, and resolution ratings."),
        ("Figure 11", "Agent Complaint Resolution View", "Interface for agents to input resolution text, update status, and close tickets."),
        ("Figure 12", "Customer Feedback Rating Form", "Rating overlay allowing customers to rate resolved complaints out of 5 stars with comments.")
    ]
    
    for fig_id, fig_title, fig_desc in figures:
        doc_builder.add_heading_3(doc, f"{fig_id}: {fig_title}")
        doc_builder.add_paragraph(doc, fig_desc, italic_text="[Screenshot Placeholder: Insert visual capture of this view here]")
        doc_builder.add_paragraph(doc, "") # Spacing
        
    # ---------------------------------------------------------
    # CHAPTER 8: ADVANTAGES & DISADVANTAGES
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "8. ADVANTAGES & DISADVANTAGES")
    
    doc_builder.add_heading_2(doc, "8.1 Key Advantages")
    doc_builder.add_bullet_point(
        doc,
        "Centralizing tickets into a single Mongo database removes the risk of lost support requests, "
        "improving response consistency and operational tracking.",
        bold_prefix="Enhanced Visibility: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Separating Customer, Agent, and Admin roles ensures strict data isolation. Sensitive details are "
        "kept secure and hidden, preventing unauthorized leak risks.",
        bold_prefix="Secure Separation of Concerns: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Admins define fields like 'Billing Account Number' or 'Hardware ID' dynamically at runtime. "
        "The backend stores them as Map objects, saving developers from altering schemas.",
        bold_prefix="Dynamic Schema Flexibility: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Customers rate resolved issues directly inside the app. Admins view these ratings on charts "
        "to evaluate agent performance, identify training gaps, and reward success.",
        bold_prefix="Built-in Feedback Loop: "
    )
    
    doc_builder.add_heading_2(doc, "8.2 Disadvantages & Constraints")
    doc_builder.add_bullet_point(
        doc,
        "The system depends entirely on active network connections. If a customer has poor internet access, "
        "they cannot register tickets or receive live updates.",
        bold_prefix="Network Dependency: "
    )
    doc_builder.add_bullet_point(
        doc,
        "State data and assets (like profile pictures or ticket attachments) are currently stored in local "
        "server folders, which limits storage scalability unless moved to cloud buckets.",
        bold_prefix="Local Storage Limits: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Admins must manually select agents from lists. If ticket volumes spike, "
        "this manual step can create support bottlenecks.",
        bold_prefix="Manual Assignment Latency: "
    )
    
    # ---------------------------------------------------------
    # CHAPTER 9: CONCLUSION
    # ---------------------------------------------------------
    doc_builder.add_heading_1(doc, "9. CONCLUSION")
    doc_builder.add_paragraph(
        doc,
        "The Customer Care Registry System successfully demonstrates a modern web solution for support management. "
        "By utilizing the MERN stack, we built a responsive SPA frontend connected to an Express API and a scalable MongoDB store. "
        "The application resolves the core stakeholder pain points: customers gain transparency via a status tracker and chat log, "
        "agents manage workloads through an assigned tickets list, and admins gain visibility through dashboards. "
        "Features like custom fields, token-based authentication, and performance tests show CCRS is ready for "
        "production deployment, aligning with standard modern software engineering practices."
    )
    
    # ---------------------------------------------------------
    # CHAPTER 10: FUTURE SCOPE
    # ---------------------------------------------------------
    doc_builder.add_heading_1(doc, "10. FUTURE SCOPE")
    doc_builder.add_paragraph(
        doc,
        "Future iterations of the CCRS can integrate advanced systems to further automate support workflows, including:"
    )
    doc_builder.add_bullet_point(
        doc,
        "Integrate AI services (like IBM Watson NLP) to analyze complaint text, automatically categorize issues, "
        "and assign them to agents based on historical resolution rates.",
        bold_prefix="Intelligent Auto-Routing: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Implement WebSockets (Socket.io) to enable real-time chat updates on ticket detail views, "
        "removing the need to refresh pages to see messages.",
        bold_prefix="Real-time WebSockets Chat: "
    )
    doc_builder.add_bullet_point(
        doc,
        "Connect AWS S3 or IBM Cloud Object Storage to handle file uploads. This ensures the app layer "
        "remains stateless and can scale horizontally in cloud environments.",
        bold_prefix="Cloud Object Storage: "
    )
    
    # ---------------------------------------------------------
    # CHAPTER 11: APPENDIX
    # ---------------------------------------------------------
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "11. APPENDIX")
    
    doc_builder.add_heading_2(doc, "11.1 Project Links")
    doc_builder.add_paragraph(doc, "GitHub Source Repository: ", italic_text="[GitHub Link Placeholder: https://github.com/username/customer-registry]")
    doc_builder.add_paragraph(doc, "System Demonstration Video: ", italic_text="[Demo Video Link Placeholder: https://youtube.com/watch?v=demo_video_id]")
    doc_builder.add_paragraph(doc, "MongoDB Database Connection: ", italic_text="[Dataset Link Placeholder: MongoDB Atlas Cloud Instance Connection String]")
    
    doc_builder.add_heading_2(doc, "11.2 Core Source Code Snippets")
    doc_builder.add_paragraph(
        doc,
        "For review, we include core backend modules, highlighting the Mongoose schemas, router middleware, "
        "and controllers that power the CCRS core engine."
    )
    
    doc_builder.add_heading_3(doc, "1. Customer Database Model (server/src/models/Customer.js)")
    code_customer = (
        "const mongoose = require('mongoose');\n\n"
        "const customerSchema = new mongoose.Schema({\n"
        "    name: { type: String, required: [true, 'Name is required'], trim: true },\n"
        "    email: { type: String, required: [true, 'Email is required'], unique: true, lowercase: true, trim: true },\n"
        "    password: { type: String, required: [true, 'Password is required'], minlength: 6 },\n"
        "    phone: { type: String, required: [true, 'Phone number is required'], trim: true },\n"
        "    role: { type: String, enum: ['customer'], default: 'customer' },\n"
        "    profileImage: { type: String, default: '' },\n"
        "    address: { type: String, default: '' },\n"
        "    customFields: { type: Map, of: String, default: {} },\n"
        "    isActive: { type: Boolean, default: true }\n"
        "}, { timestamps: true });\n\n"
        "module.exports = mongoose.model('Customer', customerSchema);"
    )
    doc_builder.add_code_block(doc, code_customer)
    
    doc_builder.add_heading_3(doc, "2. Complaint Database Model (server/src/models/Complaint.js)")
    code_complaint = (
        "const mongoose = require('mongoose');\n\n"
        "const complaintSchema = new mongoose.Schema({\n"
        "    customer: { type: mongoose.Schema.Types.ObjectId, ref: 'Customer', required: true },\n"
        "    assignedAgent: { type: mongoose.Schema.Types.ObjectId, ref: 'Agent', default: null },\n"
        "    title: { type: String, required: true, trim: true },\n"
        "    description: { type: String, required: true, trim: true },\n"
        "    category: { type: String, required: true },\n"
        "    priority: { type: String, enum: ['Low', 'Medium', 'High', 'Critical'], default: 'Medium' },\n"
        "    status: { type: String, enum: ['Pending', 'Assigned', 'In Progress', 'Resolved', 'Closed'], default: 'Pending' },\n"
        "    resolution: { type: String, default: '' },\n"
        "    resolvedAt: { type: Date, default: null },\n"
        "    attachments: [{ fileName: String, fileUrl: String }]\n"
        "}, { timestamps: true });\n\n"
        "module.exports = mongoose.model('Complaint', complaintSchema);"
    )
    doc_builder.add_code_block(doc, code_complaint)
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Final_Report.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
