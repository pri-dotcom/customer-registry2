import os
import sys
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("FUNCTIONAL SPECIFICATION DOCUMENT (FSD)\n\nCUSTOMER CARE REGISTRY SYSTEM")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Core Deliverable")
    
    # 1. Introduction
    doc_builder.add_heading_1(doc, "1. Introduction")
    doc_builder.add_paragraph(doc, "Project Title: ", bold_prefix="System Name: ", italic_text="Customer Care Registry System (CCRS)")
    
    p_team = doc.add_paragraph()
    p_team.add_run("Project Team Members & Roles:\n").bold = True
    p_team.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )
    
    doc_builder.add_paragraph(
        doc,
        "This Functional Specification Document (FSD) provides a technical and functional footprint of the Customer Care "
        "Registry System. It defines the software's components, operational architecture, API contract definitions, setup "
        "mechanisms, authentication schemes, testing paradigms, and user interface workflows. The CCRS provides a role-based "
        "environment for Customers, Agents, and Administrators, ensuring secure, transparent complaint registration and resolution."
    )
    
    # 2. Project Overview
    doc_builder.add_heading_1(doc, "2. Project Overview")
    doc_builder.add_paragraph(
        doc,
        "The CCRS is designed to modernize utility customer care. By replacing spreadsheets and emails, the system "
        "helps organizations organize incoming customer issues, assign them to verified agents based on departments, "
        "track ticket progression in real-time, and log performance feedback. Key features include role-based dashboards, "
        "JWT-secured authentication, dynamic custom registration fields, in-app ticket messaging, and administrative analytics charts."
    )
    
    # 3. Architecture
    doc_builder.add_heading_1(doc, "3. Architecture")
    doc_builder.add_paragraph(
        doc,
        "The system follows a standard three-tier pattern: Presentation Layer (Vite React Client), Logic Layer (Express/Node Server), "
        "and Data Layer (MongoDB Atlas NoSQL Store). Interaction is stateless, using JSON Web Tokens (JWT) for authentication."
    )
    doc_builder.add_bullet_point(doc, "Vite React Client: Serves the Single Page Application (SPA), managing rendering and local application states.", level=0)
    doc_builder.add_bullet_point(doc, "Express Server: Serves as the API Gateway, handling routing, validations, and database updates.", level=0)
    doc_builder.add_bullet_point(doc, "MongoDB Atlas: Stores documents in collections (Customers, Agents, Admins, Complaints, Feedback, Notifications, CustomFields).", level=0)
    
    # 4. Setup Instructions
    doc_builder.add_heading_1(doc, "4. Setup Instructions")
    doc_builder.add_paragraph(
        doc,
        "Prerequisites: Node.js (v18+), npm (v9+), and a MongoDB Atlas Cluster connection URI.\n\n"
        "Installation Steps:"
    )
    doc_builder.add_numbered_point(doc, "Clone the project repository and open it in a shell.", num_str="1. ")
    doc_builder.add_numbered_point(doc, "Navigate to the server directory and install backend dependencies: 'cd server && npm install'", num_str="2. ")
    doc_builder.add_numbered_point(doc, "Create a '.env' file inside the server directory with variables: PORT=5001, MONGO_URI=your_mongodb_uri, JWT_SECRET=your_secret_key", num_str="3. ")
    doc_builder.add_numbered_point(doc, "Navigate to the client directory and install frontend dependencies: 'cd ../client && npm install'", num_str="4. ")
    doc_builder.add_numbered_point(doc, "Create a '.env' file in client with: VITE_API_URL=http://localhost:5001/api", num_str="5. ")
    
    # 5. Folder Structure
    doc_builder.add_heading_1(doc, "5. Folder Structure")
    folder_tree = (
        "customer-registry/\n"
        "├── client/\n"
        "│   ├── public/\n"
        "│   ├── src/\n"
        "│   │   ├── components/      # Common UI elements (Navbar, Button, Card)\n"
        "│   │   ├── context/         # AuthContext state management\n"
        "│   │   ├── layouts/         # Layout grids for Admin, Agent, and Customer\n"
        "│   │   ├── pages/           # Pages (Login, Register, Dashboard, admin/, agent/)\n"
        "│   │   ├── services/        # Axios API configurations\n"
        "│   │   └── App.jsx          # Root routing and theme registry\n"
        "│   ├── index.html\n"
        "│   └── package.json\n"
        "└── server/\n"
        "    ├── src/\n"
        "    │   ├── config/          # MongoDB db.js connection file\n"
        "    │   ├── controllers/     # Controller logic (auth, complaint, feedback)\n"
        "    │   ├── middleware/      # Auth, role-based, and error middlewares\n"
        "    │   ├── models/          # Mongoose models schemas\n"
        "    │   ├── routes/          # Express API route modules\n"
        "    │   ├── utils/           # Token and password utilities\n"
        "    │   └── app.js           # Server application module config\n"
        "    ├── server.js            # Entry point listener\n"
        "    └── package.json"
    )
    doc_builder.add_code_block(doc, folder_tree)
    
    # 6. Running Application
    doc_builder.add_heading_1(doc, "6. Running the Application")
    doc_builder.add_paragraph(doc, "To start the backend Express server locally:")
    doc_builder.add_code_block(doc, "cd server\nnpm run dev   # Starts with nodemon at port 5001")
    doc_builder.add_paragraph(doc, "To start the frontend React client locally:")
    doc_builder.add_code_block(doc, "cd client\nnpm run dev   # Starts Vite server at http://localhost:5173")
    
    # 7. API Documentation
    doc.add_page_break()
    doc_builder.add_heading_1(doc, "7. API Documentation")
    
    doc_builder.add_heading_2(doc, "7.1 Authentication Endpoints")
    
    doc_builder.add_heading_3(doc, "POST /api/auth/register")
    doc_builder.add_paragraph(doc, "Registers a new user (Customer, Agent, or Admin) inside the database.")
    doc_builder.add_paragraph(doc, "Request Payload:")
    doc_builder.add_code_block(doc, '{\n  "name": "Jane Doe",\n  "email": "jane@example.com",\n  "password": "securepassword123",\n  "phone": "9876543210",\n  "role": "customer"\n}')
    doc_builder.add_paragraph(doc, "Response Body (201 Created):")
    doc_builder.add_code_block(doc, '{\n  "success": true,\n  "message": "Registration successful.",\n  "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n  "user": { "id": "60d5ec4b8f", "name": "Jane Doe", "email": "jane@example.com", "role": "customer" }\n}')
    
    doc_builder.add_heading_3(doc, "POST /api/auth/login")
    doc_builder.add_paragraph(doc, "Logs in a Customer or Administrator, returning a session token.")
    doc_builder.add_paragraph(doc, "Request Payload:")
    doc_builder.add_code_block(doc, '{\n  "email": "jane@example.com",\n  "password": "securepassword123"\n}')
    doc_builder.add_paragraph(doc, "Response Body (200 OK):")
    doc_builder.add_code_block(doc, '{\n  "success": true,\n  "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n  "user": { "id": "60d5ec4b8f", "name": "Jane Doe", "email": "jane@example.com", "role": "customer" }\n}')
    
    doc_builder.add_heading_2(doc, "7.2 Complaint Endpoints")
    
    doc_builder.add_heading_3(doc, "POST /api/complaints")
    doc_builder.add_paragraph(doc, "Creates a new customer complaint. Requires JWT Authorization Header.")
    doc_builder.add_paragraph(doc, "Request Payload:")
    doc_builder.add_code_block(doc, '{\n  "title": "Home Router Malfunction",\n  "description": "The green status LED is off and internet link is broken.",\n  "category": "Broadband",\n  "priority": "High"\n}')
    doc_builder.add_paragraph(doc, "Response Body (201 Created):")
    doc_builder.add_code_block(doc, '{\n  "success": true,\n  "complaint": {\n    "_id": "60d5ed128f",\n    "customer": "60d5ec4b8f",\n    "title": "Home Router Malfunction",\n    "description": "The green status LED is off and internet link is broken.",\n    "category": "Broadband",\n    "priority": "High",\n    "status": "Pending",\n    "createdAt": "2026-07-18T15:20:00.000Z"\n  }\n}')
    
    doc_builder.add_heading_3(doc, "GET /api/complaints")
    doc_builder.add_paragraph(doc, "Retrieves a list of complaints. For customers, returns their own complaints. For agents/admins, returns filtered queues.")
    doc_builder.add_paragraph(doc, "Response Body (200 OK):")
    doc_builder.add_code_block(doc, '{\n  "success": true,\n  "complaints": [\n    { "_id": "60d5ed128f", "title": "Home Router Malfunction", "status": "Pending", "priority": "High" }\n  ]\n}')
    
    doc_builder.add_heading_2(doc, "7.3 Feedback Endpoints")
    
    doc_builder.add_heading_3(doc, "POST /api/feedback")
    doc_builder.add_paragraph(doc, "Submits customer feedback for a resolved complaint.")
    doc_builder.add_paragraph(doc, "Request Payload:")
    doc_builder.add_code_block(doc, '{\n  "complaint": "60d5ed128f",\n  "rating": 5,\n  "comment": "Superb resolution! The agent fixed the issue in 2 hours."\n}')
    doc_builder.add_paragraph(doc, "Response Body (201 Created):")
    doc_builder.add_code_block(doc, '{\n  "success": true,\n  "feedback": { "_id": "60d5ef9a8f", "complaint": "60d5ed128f", "rating": 5, "comment": "Superb resolution!" }\n}')
    
    # 8. Authentication
    doc_builder.add_heading_1(doc, "8. Authentication")
    doc_builder.add_paragraph(
        doc,
        "Security is implemented at both transaction and storage levels. "
        "User passwords are encrypted before storage using bcrypt with a salt factor of 10. "
        "After a successful login, the server generates a JSON Web Token (JWT) signed with a secure HS256 key. "
        "The token's payload contains the user's database ID and role, which the server uses to authorize requests."
    )
    doc_builder.add_paragraph(
        doc,
        "The authMiddleware intercepts requests to secured endpoints, extracts the token from the Authorization header, "
        "verifies its signature, and appends the user details to the request object. Sub-routers then enforce role-based rules, "
        "ensuring that administrative operations (like creating custom fields) are restricted only to admin users."
    )
    
    # 9. User Interface
    doc_builder.add_heading_1(doc, "9. User Interface")
    doc_builder.add_paragraph(
        doc,
        "The web user interface is responsive, using modern UI design principles. "
        "Navigation is handled on the client-side via React Router. The primary page views are:"
    )
    doc_builder.add_bullet_point(doc, "Dashboard: Renders user profile information and summaries of active tickets.", bold_prefix="Customer Dashboard: ")
    doc_builder.add_bullet_point(doc, "Raise Complaint Form: Form with interactive validation and custom fields.", bold_prefix="Complaint Form: ")
    doc_builder.add_bullet_point(doc, "Admin Panel: Management view with list views, assignment dropdowns, and charts.", bold_prefix="Admin Panel: ")
    doc_builder.add_bullet_point(doc, "Agent Portal: Display of assigned complaints with communication timelines.", bold_prefix="Agent Portal: ")
    
    # 10. Testing
    doc_builder.add_heading_1(doc, "10. Testing")
    doc_builder.add_paragraph(
        doc,
        "The testing strategy covers Unit, Integration, and User Acceptance testing. "
        "Jest is used to run unit tests for database controllers, mock database models, and verify correct response codes. "
        "Supertest is used to run integration tests, confirming that routes are protected by the auth middleware."
    )
    
    # 11. Screenshots or Demo
    doc_builder.add_heading_1(doc, "11. Screenshots or Demo")
    doc_builder.add_paragraph(doc, "GitHub Project Repository: ", italic_text="[Insert Repository Link here]")
    doc_builder.add_paragraph(doc, "Video Demonstration Link: ", italic_text="[Insert Video URL here]")
    
    # 12. Known Issues
    doc_builder.add_heading_1(doc, "12. Known Issues")
    doc_builder.add_bullet_point(doc, "Attachments are uploaded directly to the local server disk, which prevents horizontal scaling of app servers.", bold_prefix="Local File Storage: ")
    doc_builder.add_bullet_point(doc, "No active network link leads to blank screens; the UI lacks a connection retry warning.", bold_prefix="Offline Handling: ")
    
    # 13. Future Enhancements
    doc_builder.add_heading_1(doc, "13. Future Enhancements")
    doc_builder.add_bullet_point(doc, "Move attachments to cloud buckets like AWS S3 to make the server stateless.", bold_prefix="S3 Storage Migration: ")
    doc_builder.add_bullet_point(doc, "Add Socket.io to support real-time chat updates on ticket pages without manual refresh.", bold_prefix="WebSockets Integration: ")
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "FSD.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
