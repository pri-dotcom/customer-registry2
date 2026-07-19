import os
import sys
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure the scratch directory containing doc_builder.py is in the python path
sys.path.append(r'C:\Users\priya\.gemini\antigravity-ide\brain\ed12a619-d5fe-4199-9cee-db547cd154c7\scratch')
import doc_builder

def generate():
    doc = doc_builder.create_document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(40)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("SCREENSHOTS OUTLINE GUIDE\n\nREQUIRED SYSTEM FIGURES & CAPTIONS")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor.from_string(doc_builder.PRIMARY_COLOR_HEX)
    
    doc.add_paragraph()
    doc_builder.add_cover_table(doc, date_str="18 July 2026", marks_str="Screenshots Guide - 2 Marks")

    p_team_section = doc.add_paragraph()
    p_team_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_team_section.paragraph_format.space_before = Pt(12)
    p_team_section.add_run("Project Team Members & Roles:\n").bold = True
    p_team_section.add_run(
        "• Priyanka Degala - Team Lead (PD) - Frontend Developer\n"
        "• Chandrakala Marri - Member (CM) - Backend Design & QA Testing\n"
        "• Eswarasai Marre - Member (EM) - Backend Developer\n"
        "• Triloka Pavani Mandhapati - Member (TPM) - Database Developer\n"
        "• Santhosh Nikhil Moka - Member (SNM) - Fullstack Feature Developer"
    )

    
    # Intro
    doc_builder.add_heading_1(doc, "1. Required Portal Screenshots")
    doc_builder.add_paragraph(
        doc,
        "This guide maps out the 12 screenshots required for the project report, detailing the specific "
        "UI elements, buttons, and layouts that should be captured in each."
    )
    
    headers_shots = ["Figure Number", "Screenshot Title", "UI Components to Capture", "User Role / View", "Page Guidelines"]
    data_shots = [
        ["Figure 1", "Home Landing Page", "Navigation bar with Home, Login, and Register links; project description panel; hero banner.", "Public View", "Ensure clean header layout is visible."],
        ["Figure 2", "Customer Registration Form", "Form fields (Name, Email, Password, Confirm Password, Phone); signup button.", "Public View", "Include red asterisk indicators for required fields."],
        ["Figure 3", "Customer Login Interface", "Email and password inputs; toggle to reveal password; submit button; link to register page.", "Public View", "Verify form validation errors are visible on incorrect input."],
        ["Figure 4", "Customer Dashboard Overview", "Customer name; stats cards (active, pending, resolved tickets); list of registered complaints.", "Customer View", "Dashboard list should contain at least two sample complaints."],
        ["Figure 5", "Raise Complaint Form", "Inputs for Title, Description, Category, Priority; File attachment drag-and-drop; submit button.", "Customer View", "Form should display dynamic custom fields if configured by admin."],
        ["Figure 6", "Customer Ticket Details & Chat Log", "Ticket details (title, ID, category); visual timeline tracking status; message box.", "Customer View", "Display at least one agent response in the chat log."],
        ["Figure 7", "Admin Dashboard Analytics", "Stats cards (pending, resolved, closed tickets); category charts and response latency metrics.", "Admin View", "Show representative data on charts using Chart.js."],
        ["Figure 8", "Admin User Management Panel", "Customer list (name, email, phone, status); Agent list; activate/deactivate toggles.", "Admin View", "Include a search bar and page navigation buttons."],
        ["Figure 9", "Admin Custom Fields Config", "Fields list; Add Custom Field form (name, type: text/number/date, isRequired toggle); delete button.", "Admin View", "Show at least one active custom field in the table."],
        ["Figure 10", "Agent Dashboard Overview", "Agent stats (assigned, in-progress, resolved tickets); performance rating indicators.", "Agent View", "Show active assigned ticket count card."],
        ["Figure 11", "Agent Complaint Resolution View", "Resolution notes text area; status change dropdown (In Progress, Resolved); submit button.", "Agent View", "Ensure text box contains sample resolution instructions."],
        ["Figure 12", "Customer Feedback Rating Form", "5-star rating toggle overlay; text comments input; submit feedback button.", "Customer View", "Show 5 stars selected, with sample comments entered."]
    ]
    doc_builder.add_styled_table(doc, headers_shots, data_shots, [1.0, 1.3, 2.2, 1.2, 1.3])
    
    # Save document
    out_dir = r"c:\Users\priya\Downloads\customer-registry-main (1)\project_documentation"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Screenshots_Outline.docx")
    doc.save(out_path)
    print("Successfully generated:", out_path)

if __name__ == "__main__":
    generate()
